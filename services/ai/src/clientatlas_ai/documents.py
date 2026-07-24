from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from clientatlas_ai.errors import SafeServiceError

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
SUPPORTED_TYPES = {".pdf": PDF_MIME, ".docx": DOCX_MIME}
MAX_PDF_PAGES = 500
MAX_ZIP_ENTRIES = 2_000
MAX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
PARSER_VERSION = "clientatlas-parser-1"


@dataclass(frozen=True, slots=True)
class ParsedBlock:
    text: str
    locator: dict[str, str | int]


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    blocks: tuple[ParsedBlock, ...]
    page_count: int | None


def validate_document(
    filename: str,
    declared_mime: str | None,
    content: bytes,
    *,
    max_bytes: int,
) -> str:
    safe_name = Path(filename).name
    extension = Path(safe_name).suffix.lower()
    expected_mime = SUPPORTED_TYPES.get(extension)
    if expected_mime is None:
        raise SafeServiceError("unsupported_file_extension")
    if declared_mime != expected_mime:
        raise SafeServiceError("mime_extension_mismatch")
    if not content:
        raise SafeServiceError("empty_file")
    if len(content) > max_bytes:
        raise SafeServiceError("file_too_large", status_code=413)

    if extension == ".pdf":
        if not content.startswith(b"%PDF-"):
            raise SafeServiceError("file_signature_mismatch")
    else:
        _validate_docx_container(content)
    return expected_mime


def _validate_docx_container(content: bytes) -> None:
    if not content.startswith(b"PK"):
        raise SafeServiceError("file_signature_mismatch")
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_ZIP_ENTRIES:
                raise SafeServiceError("archive_entry_limit")
            if any(entry.flag_bits & 0x1 for entry in entries):
                raise SafeServiceError("encrypted_document")
            total_size = sum(entry.file_size for entry in entries)
            if total_size > MAX_UNCOMPRESSED_BYTES:
                raise SafeServiceError("archive_expansion_limit")
            names = {entry.filename for entry in entries}
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise SafeServiceError("invalid_docx_container")
    except zipfile.BadZipFile as error:
        raise SafeServiceError("invalid_docx_container") from error


def parse_document(mime_type: str, content: bytes) -> ParsedDocument:
    if mime_type == PDF_MIME:
        return _parse_pdf(content)
    if mime_type == DOCX_MIME:
        return _parse_docx(content)
    raise SafeServiceError("unsupported_mime_type")


def _normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _parse_pdf(content: bytes) -> ParsedDocument:
    try:
        reader = PdfReader(io.BytesIO(content), strict=True)
        if reader.is_encrypted:
            raise SafeServiceError("encrypted_document")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise SafeServiceError("page_limit")
        blocks: list[ParsedBlock] = []
        for index, page in enumerate(reader.pages, start=1):
            text = _normalize_text(page.extract_text() or "")
            if text:
                blocks.append(ParsedBlock(text=text, locator={"page": index}))
        if not blocks:
            raise SafeServiceError("no_extractable_text")
        return ParsedDocument(blocks=tuple(blocks), page_count=len(reader.pages))
    except SafeServiceError:
        raise
    except Exception as error:
        raise SafeServiceError("pdf_parse_failed") from error


def _parse_docx(content: bytes) -> ParsedDocument:
    try:
        document = Document(io.BytesIO(content))
        blocks = tuple(
            ParsedBlock(
                text=text,
                locator={"paragraph": index},
            )
            for index, paragraph in enumerate(document.paragraphs, start=1)
            if (text := _normalize_text(paragraph.text))
        )
        if not blocks:
            raise SafeServiceError("no_extractable_text")
        return ParsedDocument(blocks=blocks, page_count=None)
    except SafeServiceError:
        raise
    except Exception as error:
        raise SafeServiceError("docx_parse_failed") from error


def locator_value(locator: dict[str, Any]) -> dict[str, str | int]:
    return {
        str(key): value
        for key, value in locator.items()
        if isinstance(value, (str, int))
    }
